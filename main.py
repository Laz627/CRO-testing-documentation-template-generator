
import os
import io
import math
import json
from datetime import datetime, date
from typing import List, Dict, Any, Optional

import streamlit as st
import pandas as pd

# Optional dependencies
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Document = None

APP_TITLE = "Pella CRO Self‑Service Test Builder"
VERSION = "v1.0.1"

PRIMARY_KPIS = [
    "RTA Form Completes (Lead Form Submits)",
    "RTA Form Starts (Lead Form Starts)",
    "CTA Click-through Rate",
    "Page-to-Form CTR",
    "Bounce Rate",
    "Time to First Input",
    "Scroll Depth 50%+",
    "Phone Clicks",
]

SECONDARY_KPIS = [
    "RTA Form Completes (Lead Form Submits)",
    "RTA Form Starts (Lead Form Starts)",
    "CTA Click-through Rate",
    "Page-to-Form CTR",
    "Bounce Rate",
    "Time to First Input",
    "Scroll Depth 50%+",
    "Phone Clicks",
    "Pageviews",
    "Avg. Time on Page",
]

TEST_TYPES = [
    "Copy & Messaging",
    "CTA Styling",
    "Form Friction",
    "Layout / Information Hierarchy",
    "Social Proof",
    "Urgency / Scarcity",
    "Media (Images/Video)",
    "Navigation / IA",
    "Trust & Security",
    "Performance / UX",
]

FUNNEL_STAGES = ["Awareness", "Consideration", "Conversion", "Post-Conversion"]
STATUS_OPTIONS = ["Proposed", "Planned", "In QA", "Live", "Completed", "Paused", "Discarded"]
EASIER_COMPONENTS = ["Evidence", "Alignment", "Speed", "Impact", "Effort", "Reality"]
EASIER_HELP = {
    "Evidence": "How strong is the evidence (quant + qual) that this will work? (1=weak, 5=strong)",
    "Alignment": "How aligned is this test to business goals/OKRs? (1=low, 5=high)",
    "Speed": "How quickly can we launch this? (1=slow, 5=fast)",
    "Impact": "Expected impact on primary KPI if it wins. (1=low, 5=high)",
    "Effort": "How easy is it to implement? (1=hard, 5=easy)",
    "Reality": "How realistic are assumptions & constraints? (1=weak, 5=strong)",
}



def init_session():
    if "catalog" not in st.session_state:
        st.session_state.catalog = []
    if "idea_builder_output" not in st.session_state:
        st.session_state.idea_builder_output = None
    if "selected_row_idx" not in st.session_state:
        st.session_state.selected_row_idx = None

def easier_total(easier: Dict[str, int]) -> int:
    return sum(int(easier.get(k, 0)) for k in EASIER_COMPONENTS)

def mde_from_relative(baseline_cr: float, relative_lift_pct: float) -> float:
    return baseline_cr * (relative_lift_pct / 100.0)

def sample_size_per_variant(baseline_cr: float, mde_abs: float, alpha: float = 0.05, power: float = 0.8) -> Optional[int]:
    try:
        from statistics import NormalDist
        z_alpha_2 = NormalDist().inv_cdf(1 - alpha/2)
        z_beta = NormalDist().inv_cdf(power)
        p1 = baseline_cr
        p2 = baseline_cr + mde_abs
        if p2 < 0 or p2 > 1 or p1 <= 0 or p1 >= 1 or mde_abs <= 0:
            return None
        p_bar = (p1 + p2) / 2.0
        num = (z_alpha_2 * math.sqrt(2 * p_bar * (1 - p_bar)) + z_beta * math.sqrt(p1*(1-p1) + p2*(1-p2)))
        n = (num ** 2) / ((p2 - p1) ** 2)
        return int(math.ceil(n))
    except Exception:
        return None

def estimate_duration_weeks(weekly_sessions: int, variants: int, sample_per_variant: int, allocation: float = 1.0) -> Optional[float]:
    try:
        if weekly_sessions <= 0 or variants <= 0 or sample_per_variant <= 0:
            return None
        total_required = sample_per_variant * variants
        effective_weekly = weekly_sessions * allocation
        if effective_weekly <= 0:
            return None
        weeks = total_required / effective_weekly
        return round(weeks, 2)
    except Exception:
        return None

def to_catalog_row(form_data: Dict[str, Any]) -> Dict[str, Any]:
    easier_sum = easier_total(form_data.get("easier", {}))
    return {
        "Test ID": form_data.get("test_id", ""),
        "Test Name": form_data.get("test_name", ""),
        "Hypothesis": form_data.get("hypothesis", ""),
        "Primary KPI": form_data.get("primary_kpi", ""),
        "Secondary KPIs": ", ".join(form_data.get("secondary_kpis", [])),
        "Baseline Traffic (Weekly)": form_data.get("weekly_sessions") or "",
        "Current Conversion Rate": form_data.get("baseline_cr") or "",
        "Baseline Sample Size": form_data.get("baseline_sample_size") or "",
        "Audience / Placement": form_data.get("audience_placement", ""),
        "EASIER - Evidence": form_data.get("easier", {}).get("Evidence", ""),
        "EASIER - Alignment": form_data.get("easier", {}).get("Alignment", ""),
        "EASIER - Speed": form_data.get("easier", {}).get("Speed", ""),
        "EASIER - Impact": form_data.get("easier", {}).get("Impact", ""),
        "EASIER - Effort": form_data.get("easier", {}).get("Effort", ""),
        "EASIER - Reality": form_data.get("easier", {}).get("Reality", ""),
        "EASIER Total": easier_sum,
        "Notes": form_data.get("notes", ""),
        "Test Type": ", ".join(form_data.get("test_types", [])),
        "Funnel Stage": form_data.get("funnel_stage", ""),
        "Owner": form_data.get("owner", ""),
        "Status": form_data.get("status", ""),
        "Start": form_data.get("start_date", ""),
        "End": form_data.get("end_date", ""),
        "Page URL(s)": form_data.get("page_urls", ""),
        "Variant Summary": form_data.get("variant_summary", ""),
    }

def export_word_doc(form_data: Dict[str, Any]) -> Optional[bytes]:
    if Document is None:
        return None
    doc = Document()
    title = doc.add_heading(form_data.get("test_name", "CRO Test Specification"), level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p = doc.add_paragraph()
    p.add_run("Test ID: ").bold = True; p.add_run(form_data.get("test_id", ""))
    p.add_run("\nOwner: ").bold = True; p.add_run(form_data.get("owner", ""))
    p.add_run("\nStatus: ").bold = True; p.add_run(form_data.get("status", ""))
    p.add_run("\nDate Created: ").bold = True; p.add_run(datetime.now().strftime("%Y-%m-%d"))
    p.add_run("\nPage URL(s): ").bold = True; p.add_run(form_data.get("page_urls", ""))
    doc.add_heading("Hypothesis", level=2); doc.add_paragraph(form_data.get("hypothesis", ""))
    doc.add_heading("KPIs", level=2)
    doc.add_paragraph(f"Primary KPI: {form_data.get('primary_kpi', '')}")
    doc.add_paragraph(f"Secondary/Tertiary KPIs: {', '.join(form_data.get('secondary_kpis', []))}")
    doc.add_heading("Baselines", level=2)
    doc.add_paragraph(f"Weekly Sessions (est.): {form_data.get('weekly_sessions', '')}")
    doc.add_paragraph(f"Current Conversion Rate: {form_data.get('baseline_cr', '')}")
    doc.add_paragraph(f"Baseline Sample Size (if known): {form_data.get('baseline_sample_size', '')}")
    doc.add_heading("Audience / Placement", level=2); doc.add_paragraph(form_data.get("audience_placement", ""))
    doc.add_heading("Variant Summary", level=2); doc.add_paragraph(form_data.get("variant_summary", ""))
    doc.add_heading("EASIER Scoring", level=2)
    tbl = doc.add_table(rows=2, cols=6)
    hdr_cells = tbl.rows[0].cells
    for i, comp in enumerate(EASIER_COMPONENTS): hdr_cells[i].text = comp
    val_cells = tbl.rows[1].cells
    for i, comp in enumerate(EASIER_COMPONENTS): val_cells[i].text = str(form_data.get("easier", {}).get(comp, ""))
    doc.add_paragraph(f"Total: {easier_total(form_data.get('easier', {}))}")
    doc.add_heading("Ready for Dev Checklist", level=2)
    checks = form_data.get("ready_checks", {})
    for label, done in checks.items(): doc.add_paragraph(f"[{'x' if done else ' '}] {label}")
    doc.add_heading("Notes / Risks / Dependencies", level=2); doc.add_paragraph(form_data.get("notes", ""))
    if form_data.get("power", {}):
        doc.add_heading("Power & Duration Estimate", level=2)
        power = form_data["power"]
        for k, v in power.items(): doc.add_paragraph(f"{k}: {v}")
    doc.add_paragraph("Generated by Pella CRO Self‑Service Test Builder")
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

def add_to_catalog(form_data: Dict[str, Any]):
    row = to_catalog_row(form_data)
    st.session_state.catalog.append(row)

def catalog_df() -> pd.DataFrame:
    if not st.session_state.catalog:
        return pd.DataFrame(columns=[
            "Test ID","Test Name","Hypothesis","Primary KPI","Secondary KPIs",
            "Baseline Traffic (Weekly)","Current Conversion Rate","Baseline Sample Size",
            "Audience / Placement","EASIER - Evidence","EASIER - Alignment","EASIER - Speed",
            "EASIER - Impact","EASIER - Effort","EASIER - Reality","EASIER Total","Notes",
            "Test Type","Funnel Stage","Owner","Status","Start","End","Page URL(s)","Variant Summary"
        ])
    return pd.DataFrame(st.session_state.catalog)

def export_catalog_xlsx(df: pd.DataFrame) -> bytes:
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="Test Log")
    return output.getvalue()

def header():
    st.title(APP_TITLE)
    st.caption(f"{VERSION} — Guided test ideation, spec, validation, and exports (Word & Excel).")



def sidebar_nav() -> str:
    page = st.sidebar.radio("Navigate", ["Create Test", "Validate", "Catalog & Export", "Admin"], index=0)
    st.sidebar.markdown("---")
    st.sidebar.caption(f"{APP_TITLE} — {VERSION}")
    return page

def idea_builder() -> Optional[Dict[str, Any]]:
    st.subheader("Phase 1 — Idea Builder (Optional)")
    st.caption("Describe your goal and page context. I’ll propose test ideas, hypotheses, and suggested KPIs.")
    goal = st.text_input("Goal (e.g., Increase form completion by 10%)")
    page_url = st.text_input("Primary Page URL (optional)")
    test_type = st.selectbox("Test Type (to guide suggestions)", TEST_TYPES, index=2)
    colA, colB = st.columns(2)
        generated = None
    if st.button("Generate Ideas", use_container_width=True):
        generated = [
            {"test_name": f"{test_type}: Variant Messaging Emphasis",
             "hypothesis": f"By refining {test_type.lower()} to better align with user intent for {page_url or 'the selected page'}, we expect an uplift in the primary KPI.",
             "suggested_primary_kpi": "RTA Form Completes (Lead Form Submits)",
             "suggested_secondary_kpis": ["RTA Form Starts (Lead Form Starts)", "CTA Click-through Rate"]},
            {"test_name": f"{test_type}: Reduce Friction Element",
             "hypothesis": f"By removing/reducing a friction element tied to {test_type.lower()}, users will progress more reliably to the form and complete it.",
             "suggested_primary_kpi": "RTA Form Completes (Lead Form Submits)",
             "suggested_secondary_kpis": ["Page-to-Form CTR", "Time to First Input"]},
            {"test_name": f"{test_type}: Trust & Proof",
             "hypothesis": f"By adding clear trust signals relevant to {test_type.lower()}, we reduce uncertainty and increase completions.",
             "suggested_primary_kpi": "RTA Form Completes (Lead Form Submits)",
             "suggested_secondary_kpis": ["RTA Form Starts (Lead Form Starts)", "Bounce Rate"]},
        ]
        if False: # OpenAI removed
            try:
                # Support OpenAI v1.x client if available; otherwise fallback to legacy lib if present
                pass
        st.session_state.idea_builder_output = generated
    out = st.session_state.idea_builder_output
    if out:
        st.write("### Candidate Ideas")
        for i, idea in enumerate(out, 1):
            with st.expander(f"Idea {i}: {idea.get('test_name','')}"):
                st.write("**Hypothesis:**", idea.get("hypothesis",""))
                st.write("**Suggested Primary KPI:**", idea.get("suggested_primary_kpi",""))
                st.write("**Suggested Secondary KPIs:**", ", ".join(idea.get("suggested_secondary_kpis", [])))
                if st.button(f"Use this idea (#{i})", key=f"use_{i}"):
                    return idea
    return None

def create_test_form(prefill: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    st.subheader("Phase 2 — Spec & Validate")
    col1, col2 = st.columns(2)
    with col1: test_id = st.text_input("Test ID (user-entered)", value=prefill.get("test_id","") if prefill else "")
    with col2: test_name = st.text_input("Test Name", value=prefill.get("test_name","") if prefill else "")
    hypothesis = st.text_area("Hypothesis", height=120, value=prefill.get("hypothesis","") if prefill else "")
    col3, col4 = st.columns(2)
    with col3:
        primary_kpi = st.selectbox("Primary KPI", PRIMARY_KPIS,
                                   index=PRIMARY_KPIS.index(prefill.get("suggested_primary_kpi")) if prefill and prefill.get("suggested_primary_kpi") in PRIMARY_KPIS else 0)
    with col4:
        secondary_kpis = st.multiselect("Secondary/Tertiary KPIs", SECONDARY_KPIS,
                                        default=prefill.get("suggested_secondary_kpis", []) if prefill else [])
    st.markdown("#### Baselines")
    col5, col6, col7 = st.columns(3)
    with col5: weekly_sessions = st.number_input("Weekly Sessions (est.)", min_value=0, value=0, step=100)
    with col6: baseline_cr = st.number_input("Current Conversion Rate (0.0–1.0)", min_value=0.0, max_value=1.0, value=0.0, step=0.001, format="%.3f")
    with col7: baseline_sample_size = st.number_input("Baseline Sample Size (if known)", min_value=0, value=0, step=100)
    st.markdown("#### Audience / Placement")
    funnel_stage = st.selectbox("Funnel Stage", FUNNEL_STAGES, index=2)
    page_urls = st.text_input("Page URL(s) (comma-separated)")
    audience_placement = st.text_area("Audience / Placement details", height=80)
    test_types = st.multiselect("Test Type(s)", TEST_TYPES, default=[prefill.get("test_type")] if prefill and prefill.get("test_type") in TEST_TYPES else [])
    variant_summary = st.text_area("Variant Summary (Control vs Variant(s) - what changes?)", height=100)
    st.markdown("#### EASIER Scoring (1–5 each, equal weights)")
    easier_scores = {}; cols = st.columns(6)
    for idx, comp in enumerate(EASIER_COMPONENTS):
        with cols[idx]:
            easier_scores[comp] = st.slider(comp, 1, 5, 3, help=EASIER_HELP[comp])
    st.info(f"EASIER Total: **{easier_total(easier_scores)}** (higher = higher priority)")
    notes = st.text_area("Notes / Risks / Dependencies", height=120)
    st.markdown("#### Governance & Tracking")
    col8, col9, col10 = st.columns(3)
    with col8: owner = st.text_input("Owner (Requester)")
    with col9: status = st.selectbox("Status", STATUS_OPTIONS, index=0)
    with col10: ready_for_dev = st.checkbox("Mark Ready for Dev", value=False)
    col11, col12 = st.columns(2)
    with col11: start_date = st.date_input("Start Date", value=None)
    with col12: end_date = st.date_input("End Date", value=None)
    st.markdown("#### Ready for Dev Checklist")
    check_labels = [
        "Analytics events are defined (Primary + Secondary KPIs)",
        "Experiment scoping and exclusion rules defined",
        "QA plan documented",
        "Variant specs / mockups attached or linked",
        "Rollout/rollback plan defined",
    ]
    ready_checks = {}; c1, c2 = st.columns(2)
    for i, label in enumerate(check_labels):
        with (c1 if i % 2 == 0 else c2):
            ready_checks[label] = st.checkbox(label, value=False)
    data = {
        "test_id": test_id.strip(),
        "test_name": test_name.strip(),
        "hypothesis": hypothesis.strip(),
        "primary_kpi": primary_kpi,
        "secondary_kpis": secondary_kpis,
        "weekly_sessions": int(weekly_sessions) if weekly_sessions else 0,
        "baseline_cr": float(baseline_cr) if baseline_cr else 0.0,
        "baseline_sample_size": int(baseline_sample_size) if baseline_sample_size else 0,
        "funnel_stage": funnel_stage,
        "page_urls": page_urls.strip(),
        "audience_placement": audience_placement.strip(),
        "test_types": test_types,
        "variant_summary": variant_summary.strip(),
        "easier": easier_scores,
        "notes": notes.strip(),
        "owner": owner.strip(),
        "status": status,
        "ready_for_dev": ready_for_dev,
        "ready_checks": ready_checks,
        "start_date": start_date.isoformat() if isinstance(start_date, date) else "",
        "end_date": end_date.isoformat() if isinstance(end_date, date) else "",
    }
    return data

def validate_block(form_data: Dict[str, Any]) -> Dict[str, Any]:
    st.subheader("Power & Duration Estimator")
    col1, col2, col3 = st.columns(3)
    with col1: variants = st.number_input("Number of Variants (incl. Control)", min_value=2, value=2, step=1)
    with col2: alpha = st.number_input("Alpha (Type I error)", min_value=0.001, max_value=0.2, value=0.05, step=0.001, format="%.3f")
    with col3: power = st.number_input("Desired Power", min_value=0.5, max_value=0.99, value=0.8, step=0.01, format="%.2f")
    col4, col5 = st.columns(2)
    with col4: mde_input_mode = st.radio("MDE Input", ["Relative % lift", "Absolute (delta in CR)"], index=0, horizontal=True)
    with col5: allocation = st.slider("Traffic Allocation to Experiment", 0.1, 1.0, 1.0, step=0.05)
    baseline_cr = form_data.get("baseline_cr", 0.0); weekly_sessions = form_data.get("weekly_sessions", 0)
    if mde_input_mode == "Relative % lift":
        rel = st.number_input("Relative Lift % (e.g., 10 = +10%)", min_value=0.1, value=10.0, step=0.1)
        mde_abs = mde_from_relative(baseline_cr, rel)
    else:
        mde_abs = st.number_input("Absolute MDE (delta in CR, e.g., 0.01 = +1pp)", min_value=0.001, value=0.01, step=0.001, format="%.3f")
    n_per_variant = sample_size_per_variant(baseline_cr, mde_abs, alpha=alpha, power=power)
    weeks = estimate_duration_weeks(weekly_sessions, int(variants), n_per_variant or 0, allocation=allocation)
    results = {
        "Baseline CR": f"{baseline_cr:.3f}",
        "MDE (absolute)": f"{mde_abs:.3f}" if mde_abs is not None else "—",
        "Alpha": f"{alpha:.3f}",
        "Power": f"{power:.2f}",
        "Variants": int(variants),
        "Sample Size per Variant": f"{n_per_variant:,}" if n_per_variant else "—",
        "Est. Duration (weeks)": f"{weeks}" if weeks is not None else "—",
        "Weekly Sessions (effective)": f"{int(weekly_sessions * allocation):,}",
    }
    st.json(results); return results

def create_test_page():
    st.header("Create Test")
    idea = idea_builder()
    prefill = None
    if idea:
        prefill = {"test_name": idea.get("test_name", ""),
                   "hypothesis": idea.get("hypothesis", ""),
                   "suggested_primary_kpi": idea.get("suggested_primary_kpi", ""),
                   "suggested_secondary_kpis": idea.get("suggested_secondary_kpis", []),
                   "test_type": None}
    form_data = create_test_form(prefill=prefill)
    if st.button("Validate & Estimate Power/Duration", type="primary", use_container_width=True):
        power_results = validate_block(form_data)
        form_data["power"] = power_results
        st.session_state.last_form_data = form_data
    st.divider()
    colA, colB, colC = st.columns(3)
    with colA:
        if st.button("Add to Catalog", use_container_width=True):
            add_to_catalog(form_data); st.success("Added to catalog below.")
    with colB:
        if st.button("Download as Word (.docx)", use_container_width=True):
            content = export_word_doc(form_data)
            if content:
                st.download_button("Save .docx", data=content, file_name=f"{form_data.get('test_id','test')}.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                st.error("python-docx not available (install python-docx).")
    with colC: st.write("")
    st.markdown("### Catalog (this session)")
    df = catalog_df(); st.dataframe(df, use_container_width=True)

def validate_page():
    st.header("Validate")
    if "last_form_data" not in st.session_state:
        st.info("No form data found. Fill the form in 'Create Test' first, then click 'Validate & Estimate Power/Duration'."); return
    form_data = st.session_state.last_form_data
    st.markdown("Using last form submission:"); st.json({k: v for k, v in form_data.items() if k != "ready_checks"})
    st.divider(); st.markdown("### Re-run Power & Duration (optional)")
    power_results = validate_block(form_data); form_data["power"] = power_results

def catalog_page():
    st.header("Catalog & Export")
    df = catalog_df(); st.dataframe(df, use_container_width=True)
    col1, col2 = st.columns(2)
    with col1:
        csv_data = df.to_csv(index=False).encode("utf-8")
        st.download_button("Download CSV", data=csv_data, file_name="cro_test_log.csv", mime="text/csv", use_container_width=True)
    with col2:
        xlsx_data = export_catalog_xlsx(df)
        st.download_button("Download Excel (.xlsx)", data=xlsx_data, file_name="cro_test_log.xlsx",
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
    st.markdown("#### Export Selected Row to Word")
    idx = st.number_input("Row index (0-based)", min_value=0, max_value=max(len(df)-1, 0), value=0) if len(df) else 0
    if len(df) and st.button("Export Selected Row", use_container_width=True):
        row = df.iloc[int(idx)].to_dict()
        form_like = {
            "test_id": row.get("Test ID",""),
            "test_name": row.get("Test Name",""),
            "hypothesis": row.get("Hypothesis",""),
            "primary_kpi": row.get("Primary KPI",""),
            "secondary_kpis": [x.strip() for x in (row.get("Secondary KPIs","") or "").split(",") if x.strip()],
            "weekly_sessions": row.get("Baseline Traffic (Weekly)", ""),
            "baseline_cr": row.get("Current Conversion Rate",""),
            "baseline_sample_size": row.get("Baseline Sample Size",""),
            "audience_placement": row.get("Audience / Placement",""),
            "easier": {
                "Evidence": row.get("EASIER - Evidence",""),
                "Alignment": row.get("EASIER - Alignment",""),
                "Speed": row.get("EASIER - Speed",""),
                "Impact": row.get("EASIER - Impact",""),
                "Effort": row.get("EASIER - Effort",""),
                "Reality": row.get("EASIER - Reality",""),
            },
            "notes": row.get("Notes",""),
            "test_types": [x.strip() for x in (row.get("Test Type","") or "").split(",") if x.strip()],
            "funnel_stage": row.get("Funnel Stage",""),
            "owner": row.get("Owner",""),
            "status": row.get("Status",""),
            "start_date": row.get("Start",""),
            "end_date": row.get("End",""),
            "page_urls": row.get("Page URL(s)",""),
            "variant_summary": row.get("Variant Summary",""),
        }
        content = export_word_doc(form_like)
        if content:
            st.download_button("Download Selected as .docx", data=content, file_name=f"{row.get('Test ID','test')}.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        else:
            st.error("python-docx not available (install python-docx).")

def admin_page():
    st.header("Admin")
    st.caption("Configure dropdowns and defaults (local to your session).")
    st.write("Primary KPI options:", PRIMARY_KPIS)
    st.write("Secondary KPI options:", SECONDARY_KPIS)
    st.write("Test Types:", TEST_TYPES)
    st.write("Funnel Stages:", FUNNEL_STAGES)
    st.write("Status Options:", STATUS_OPTIONS)
    st.info("EASIER uses equal weights in this version.")
    st.markdown("---")
    st.caption("OpenAI (optional): set OPENAI_API_KEY in Streamlit secrets or environment.")

def main():
    st.set_page_config(page_title=APP_TITLE, page_icon="🧪", layout="wide")
    init_session()
    st.sidebar.markdown("## Navigation")
    page = sidebar_nav()
    st.sidebar.markdown("---")
    st.sidebar.caption(f"{APP_TITLE} — {VERSION}")
    if page == "Create Test":
        header(); create_test_page()
    elif page == "Validate":
        header(); validate_page()
    elif page == "Catalog & Export":
        header(); catalog_page()
    elif page == "Admin":
        header(); admin_page()
    else:
        header(); create_test_page()

if __name__ == "__main__":
    main()
